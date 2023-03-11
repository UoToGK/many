#!/usr/bin/env python
# coding=utf-8
"""
Author: UoToGK 
LastEditors: UoToGK
Date: 2022-04-17 17:52:42
LastEditTime: 2022-04-17 17:52:46
FilePath: 52pj_weibo_text.py
Description: 
Copyright (c) 2022 by DyTheme, All Rights Reserved. 
"""

# Standard Library
import json
import time

# Third Library
import docx
import requests

s = requests.Session()
page = 1
headers = {
    "user-agent": "'Mozilla/5.0 (iPhone; CPU iPhone OS 13_2_3 like Mac OS X) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/13.0.3 Mobile/15E148 Safari/604.1'",
}
# 日期格式化处理，微博获取到的原始时间比较拉胯
weekdict = {
    "Mon": "星期一",
    "Tue": "星期二",
    "Wed": "星期三",
    "Thu": "星期四",
    "Fri": "星期五",
    "Sat": "星期六",
    "Sun": "星期日",
}
mondict = {
    "Jan": "01",
    "Feb": "02",
    "Mar": "03",
    "Apr": "04",
    "May": "05",
    "Jun": "06",
    "Jul": "07",
    "Aug": "08",
    "Sep": "09",
    "Oct": "10",
    "Nov": "11",
    "Dec": "12",
}
file = docx.Document()  # 创建内存中的word文档对象


def getList(since_id=None):
    global page
    while True:
        url = "https://m.weibo.cn/api/container/getIndex?type=uid&[b]value=6367430139&containerid=1076036367430139&since_id={}".format(
            since_id
        )
        res = s.get(url, headers=headers)
        r = json.loads(res.text)
        since_id = r["data"]["cardlistInfo"]["since_id"]
        list = r["data"]["cards"]

        for item in list:
            if item["mblog"]["isLongText"]:
                getLongText(file, item["mblog"]["id"], item["mblog"]["created_at"])
            else:
                txt = item["mblog"]["text"].replace("<br />", "\n")
                writeFile(file, item["mblog"]["created_at"], txt)
        page = page + 1
        time.sleep(2)
        file.save("wb.docx")


# 获取全文
def getLongText(file, id, date):
    url = "https://m.weibo.cn/statuses/extend?id={}".format(id)
    res = s.get(url)
    try:
        r = json.loads(res.text)
        txt = r["data"]["longTextContent"].replace("<br />", "\n")
        writeFile(file, date, txt)
        print(r["data"]["longTextContent"])
        print("写入成功，{}".format(url))
    except:
        print("写入文件出错，跳过···{}".format(url))
    time.sleep(2)


def fotmatDate(s):
    list = s.split(" ")
    date = "{}-{}-{} {} {}".format(
        list[-1], mondict[list[1]], list[2], list[3], weekdict[list[0]]
    )
    return date


def writeFile(file, date, data):
    # with open('微博mobile.txt','a',encoding='utf-8') as f:
    #     f.write('\n\n')
    #     f.write(fotmatDate(date))
    #     f.write('\n')
    #     f.write(data)
    # 添加段落，格式是：时间-换行是正文
    file.add_heading(fotmatDate(date), level=1)
    file.add_paragraph(data)


if __name__ == "__main__":
    getList()
